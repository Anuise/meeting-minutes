"""`mm scan-docx` / `mm apply-docx` — 把 Docx Source 打洞成 Docx Template。

沿用同一形狀：準備 fixture 目錄、執行一次子指令、
斷言檔案系統的結果與 stdout 的 JSON。
"""

import json
import subprocess
import sys
from pathlib import Path

MM = Path(__file__).resolve().parent.parent / "scripts" / "mm.py"

SOURCE = "client-signoff.docx"
MEETING = "2026-07-28-client-signoff"

TITLE = "2026 年度客戶驗收會議記錄"
# Word 常把一句話拆成好幾個 run（改過字型、拼字檢查、追蹤修訂都會）。
# 這份 fixture 刻意重現它：標題三個 run，其中一段是粗體。
TITLE_RUNS = (("2026 年度", True), ("客戶驗收", False), ("會議記錄", False))
HEADER = "MoBagel 行動貝果｜內部文件"
FOOTER = "第 1 頁"
CLOSING = "本次驗收由 王小明 主持。"

MAPPING = [
    {"text": TITLE, "variable": "{{ meta.title }}"},
    {"text": "2026-07-28", "variable": "{{ meta.datetime }}"},
    {"text": "王小明", "variable": "{{ meta.chair }}"},
]

RECORD = """\
meta:
  title: 專案週會
  datetime: 2026-07-28 14:00
  chair: 李小華
"""


def run_mm(*args, stdin=None):
    return subprocess.run(
        [sys.executable, str(MM), *args],
        input=stdin,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )


def init(root):
    result = run_mm("init", "--root", str(root))
    assert result.returncode == 0, result.stderr


def write_docx_source(root, name=SOURCE):
    """一份客戶樣板該有的東西都湊齊：頁首頁尾、有樣式的表格、跨 run 的句子。"""
    from docx import Document

    target = root / "templates" / "docx-source" / name
    target.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    title = document.add_paragraph()
    for piece, bold in TITLE_RUNS:
        title.add_run(piece).bold = bold

    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    for row, (label, value) in enumerate((("日期", "2026-07-28"), ("主席", "王小明"))):
        table.rows[row].cells[0].text = label
        table.rows[row].cells[1].text = value

    document.add_paragraph(CLOSING)

    section = document.sections[0]
    section.header.paragraphs[0].text = HEADER
    section.footer.paragraphs[0].text = FOOTER

    document.save(str(target))
    return target


def texts(path):
    """把一份 .docx 拆成本文段落、表格儲存格、頁首與頁尾的文字。"""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(str(path))
    body = []
    for element in document.element.body:
        if element.tag.endswith("}p"):
            text = Paragraph(element, document).text.strip()
            if text:
                body.append(text)
        elif element.tag.endswith("}tbl"):
            for row in Table(element, document).rows:
                body.extend(cell.text.strip() for cell in row.cells)

    section = document.sections[0]
    return {
        "body": body,
        "header": section.header.paragraphs[0].text,
        "footer": section.footer.paragraphs[0].text,
        "table_style": document.tables[0].style.name if document.tables else None,
    }


def scan(root, source=SOURCE):
    result = run_mm("scan-docx", source, "--root", str(root))
    assert result.returncode == 0, result.stderr
    return json.loads(result.stdout)


def apply(root, mapping=MAPPING, source=SOURCE, *extra):
    result = run_mm(
        "apply-docx",
        source,
        "--root",
        str(root),
        *extra,
        stdin=json.dumps(mapping, ensure_ascii=False),
    )
    assert result.returncode == 0, result.stderr
    return json.loads(result.stdout)


def prepared(root):
    """已跑過 init、放好一份 Docx Source 的骨架。"""
    init(root)
    write_docx_source(root)
    return root


def template(root, name=SOURCE):
    return root / "templates" / "docx" / name


def test_scan_lists_paragraphs_cells_and_header_footer(tmp_path):
    payload = scan(prepared(tmp_path))

    items = payload["items"]
    listed = {item["text"] for item in items}
    assert TITLE in listed
    assert "2026-07-28" in listed
    assert CLOSING in listed
    assert HEADER in listed
    assert FOOTER in listed

    by_text = {item["text"]: item for item in items}
    # 儲存格帶著座標，agent 才判斷得出「左邊是標籤、右邊是值」
    assert by_text["2026-07-28"]["kind"] == "cell"
    assert by_text["2026-07-28"]["where"] == "body"
    assert (by_text["2026-07-28"]["row"], by_text["2026-07-28"]["column"]) == (0, 1)
    assert by_text[TITLE]["kind"] == "paragraph"
    assert by_text[HEADER]["where"] == "header"
    assert by_text[FOOTER]["where"] == "footer"


def test_scan_reports_a_sentence_split_across_runs_as_one_item(tmp_path):
    payload = scan(prepared(tmp_path))

    # 使用者眼裡那是一句話，就該以一句話的形式出現在對照表上
    assert [item["text"] for item in payload["items"]].count(TITLE) == 1
    assert not any(item["text"] == "客戶驗收" for item in payload["items"])


def test_scan_skips_empty_paragraphs(tmp_path):
    payload = scan(prepared(tmp_path))

    assert all(item["text"].strip() for item in payload["items"])


def test_scan_leaves_the_docx_source_untouched(tmp_path):
    source = write_docx_source(prepared(tmp_path))
    before = source.read_bytes()

    scan(tmp_path)

    assert source.read_bytes() == before


def test_missing_docx_source_is_an_error(tmp_path):
    init(tmp_path)

    result = run_mm("scan-docx", "nope.docx", "--root", str(tmp_path))

    assert result.returncode == 1
    assert "nope.docx" in result.stderr


def test_applied_template_renders_with_docxtpl(tmp_path):
    prepared(tmp_path)
    payload = apply(tmp_path)

    assert payload["docx_template"] == str(template(tmp_path))
    # 落在 templates/docx/ 才會出現在 mm list 的可用模板裡
    listing = run_mm("list", "--root", str(tmp_path))
    assert SOURCE in json.loads(listing.stdout)["docx_templates"]

    record = tmp_path / "records" / f"{MEETING}.yaml"
    record.parent.mkdir(parents=True, exist_ok=True)
    record.write_text(RECORD, encoding="utf-8")
    result = run_mm(
        "render",
        MEETING,
        "--markdown-template",
        "default.md.j2",
        "--docx-template",
        SOURCE,
        "--root",
        str(tmp_path),
    )
    assert result.returncode == 0, result.stderr

    content = texts(tmp_path / "output" / MEETING / "minutes.docx")
    assert "專案週會" in content["body"]
    assert "2026-07-28 14:00" in content["body"]
    assert "本次驗收由 李小華 主持。" in content["body"]


def test_sentence_split_across_runs_is_replaced(tmp_path):
    prepared(tmp_path)

    payload = apply(tmp_path)

    assert payload["unmatched"] == []
    body = texts(template(tmp_path))["body"]
    assert "{{ meta.title }}" in body
    assert not any(TITLE in text for text in body)


def test_the_same_text_is_replaced_everywhere_it_appears(tmp_path):
    prepared(tmp_path)

    payload = apply(tmp_path)

    # 「王小明」在儲存格與段落各出現一次，兩處都得打洞，且回報看得出打了幾個
    replaced = {entry["text"]: entry["count"] for entry in payload["replaced"]}
    assert replaced["王小明"] == 2
    body = texts(template(tmp_path))["body"]
    assert "{{ meta.chair }}" in body
    assert "本次驗收由 {{ meta.chair }} 主持。" in body


def test_merged_cells_are_visited_once(tmp_path):
    # 合併過的儲存格在同一列裡會重複出現，重複計數會讓使用者以為文件上有兩處
    from docx import Document

    init(tmp_path)
    target = tmp_path / "templates" / "docx-source" / SOURCE
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].merge(table.rows[0].cells[1]).text = "會議地點：台北辦公室"
    document.save(str(target))

    scanned = scan(tmp_path)
    payload = apply(tmp_path, [{"text": "台北辦公室", "variable": "{{ meta.location }}"}])

    assert [item["text"] for item in scanned["items"]] == ["會議地點：台北辦公室"]
    assert payload["replaced"] == [
        {"text": "台北辦公室", "variable": "{{ meta.location }}", "count": 1}
    ]


def test_a_paragraph_with_a_hyperlink_is_left_alone(tmp_path):
    """超連結的文字不在 run 裡，合併 run 搬不動它。

    硬做會把連結文字擠到別的位置、或整段抓不到 run 而爆掉。這種段落整段跳過，
    原文出現在 unmatched，使用者知道自己有一段沒被打洞。
    """
    from docx import Document
    from docx.oxml import OxmlElement

    init(tmp_path)
    document = Document()
    paragraph = document.add_paragraph("詳見 ")
    hyperlink = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "公司網站"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    paragraph.add_run(" 的 2026-07-28 公告")
    document.save(str(tmp_path / "templates" / "docx-source" / SOURCE))

    payload = apply(tmp_path, [{"text": "2026-07-28", "variable": "{{ meta.datetime }}"}])

    assert payload["unmatched"] == ["2026-07-28"]
    assert texts(template(tmp_path))["body"] == ["詳見 公司網站 的 2026-07-28 公告"]


def test_apply_leaves_the_docx_source_untouched(tmp_path):
    source = write_docx_source(prepared(tmp_path))
    before = source.read_bytes()

    apply(tmp_path)

    assert source.read_bytes() == before


def test_header_footer_and_table_style_survive_the_punching(tmp_path):
    prepared(tmp_path)

    apply(tmp_path)

    content = texts(template(tmp_path))
    assert content["header"] == HEADER
    assert content["footer"] == FOOTER
    assert content["table_style"] == "Table Grid"


def test_unmatched_mapping_entries_are_reported(tmp_path):
    prepared(tmp_path)

    payload = apply(tmp_path, MAPPING + [{"text": "不存在的字", "variable": "{{ x }}"}])

    assert payload["unmatched"] == ["不存在的字"]


def test_existing_docx_template_is_not_overwritten_without_force(tmp_path):
    prepared(tmp_path)
    apply(tmp_path)
    before = template(tmp_path).read_bytes()

    result = run_mm(
        "apply-docx",
        SOURCE,
        "--root",
        str(tmp_path),
        stdin=json.dumps([{"text": TITLE, "variable": "{{ x }}"}], ensure_ascii=False),
    )

    assert result.returncode == 1
    assert SOURCE in result.stderr
    assert template(tmp_path).read_bytes() == before

    payload = apply(tmp_path, [{"text": TITLE, "variable": "{{ x }}"}], SOURCE, "--force")
    assert payload["unmatched"] == []
    assert "{{ x }}" in texts(template(tmp_path))["body"]


def test_marker_rows_become_a_repeating_table(tmp_path):
    """打洞只換文字，變不出新的一列。

    客戶樣板的待辦事項表格只有一列範例，要讓它跟著 Minutes Record 長出多列，
    使用者得先在 Docx Source 上、範例列的前後各加一列並隨手打個標記字，
    再把那兩個標記換成 docxtpl 的 `{%tr %}`。這是 mm-template 教使用者的做法，
    所以這裡確認它真的渲得出來。
    """
    from docx import Document

    init(tmp_path)
    document = Document()
    table = document.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "事項"
    table.rows[0].cells[1].text = "負責人"
    table.rows[1].cells[0].text = "迴圈開始"
    table.rows[2].cells[0].text = "補上錯誤訊息"
    table.rows[2].cells[1].text = "張大同"
    table.rows[3].cells[0].text = "迴圈結束"
    document.save(str(tmp_path / "templates" / "docx-source" / SOURCE))

    apply(
        tmp_path,
        [
            {"text": "迴圈開始", "variable": "{%tr for item in action_items %}"},
            {"text": "迴圈結束", "variable": "{%tr endfor %}"},
            {"text": "補上錯誤訊息", "variable": "{{ item.task }}"},
            {"text": "張大同", "variable": "{{ item.owner }}"},
        ],
    )

    record = tmp_path / "records" / f"{MEETING}.yaml"
    record.parent.mkdir(parents=True, exist_ok=True)
    record.write_text(
        "action_items:\n"
        "  - { task: 補上登入頁的錯誤訊息, owner: 張大同 }\n"
        "  - { task: 確認驗收時程, owner: 李小華 }\n",
        encoding="utf-8",
    )
    result = run_mm(
        "render",
        MEETING,
        "--markdown-template",
        "default.md.j2",
        "--docx-template",
        SOURCE,
        "--root",
        str(tmp_path),
    )
    assert result.returncode == 0, result.stderr

    body = texts(tmp_path / "output" / MEETING / "minutes.docx")["body"]
    assert "補上登入頁的錯誤訊息" in body
    assert "確認驗收時程" in body
    assert "迴圈開始" not in body


def test_output_name_can_differ_from_the_source(tmp_path):
    prepared(tmp_path)

    payload = apply(tmp_path, MAPPING, SOURCE, "--output", "signoff-2026.docx")

    assert payload["docx_template"] == str(template(tmp_path, "signoff-2026.docx"))
    assert not template(tmp_path).exists()


def test_output_must_be_a_docx(tmp_path):
    prepared(tmp_path)

    result = run_mm(
        "apply-docx",
        SOURCE,
        "--output",
        "signoff",
        "--root",
        str(tmp_path),
        stdin=json.dumps(MAPPING, ensure_ascii=False),
    )

    assert result.returncode == 1
    assert not any((tmp_path / "templates" / "docx").glob("signoff*"))


def test_empty_mapping_is_an_error(tmp_path):
    prepared(tmp_path)

    result = run_mm("apply-docx", SOURCE, "--root", str(tmp_path), stdin="")

    assert result.returncode == 1
    assert not template(tmp_path).exists()


def test_malformed_mapping_is_an_error(tmp_path):
    prepared(tmp_path)

    broken = run_mm("apply-docx", SOURCE, "--root", str(tmp_path), stdin="{not json")
    assert broken.returncode == 1

    # 少了 variable 的項目不能默默跳過——使用者以為那一格打了洞
    incomplete = run_mm(
        "apply-docx",
        SOURCE,
        "--root",
        str(tmp_path),
        stdin=json.dumps([{"text": TITLE}], ensure_ascii=False),
    )
    assert incomplete.returncode == 1

    duplicated = run_mm(
        "apply-docx",
        SOURCE,
        "--root",
        str(tmp_path),
        stdin=json.dumps(
            [
                {"text": TITLE, "variable": "{{ a }}"},
                {"text": TITLE, "variable": "{{ b }}"},
            ],
            ensure_ascii=False,
        ),
    )
    assert duplicated.returncode == 1

    assert not template(tmp_path).exists()
