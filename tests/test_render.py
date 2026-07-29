"""`mm render` — Minutes Record 套 Markdown Template 變成 Deliverable。

沿用同一形狀：準備 fixture 目錄、執行一次子指令、
斷言檔案系統的結果與 stdout 的 JSON。
"""

import json
import shutil
import subprocess
import sys
from pathlib import Path

MM = Path(__file__).resolve().parent.parent / "scripts" / "mm.py"

MEETING = "2026-07-28-project-weekly"
TEMPLATE = "default.md.j2"
DOCX_TEMPLATE = "default.docx"

# 一份「有填、有空、有巢狀、有清單」都湊齊的 Minutes Record：
# 第二則議題與第二筆待辦刻意留空，用來驗證空欄位與未填變數清單。
RECORD = """\
meta:
  title: 專案週會
  datetime: 2026-07-28 14:00
  location: ''
  chair: 王小明
  recorder: 李小華
  attendees: [王小明, 李小華, 張大同]
  absentees: []

topics:
  - title: 進度回顧
    discussion: 前端已完成登入頁
    resolutions:
      - text: 登入頁下週上線
        source: notes/2026-07-28-project-weekly/slides.pptx.md#L12
      - text: 舊版並行一週
        source: notes/2026-07-28-project-weekly/slides.pptx.md#L18
  - title: 風險盤點
    discussion: ''
    resolutions: []

action_items:
  - task: 補上登入頁的錯誤訊息
    owner: 張大同
    due: 2026-08-04
    source: notes/2026-07-28-project-weekly/agenda.pdf.md#L3
  - task: 確認驗收時程
    owner: '   '
    due: ''
    source:

next_meeting: 2026-08-04 14:00
"""

# 順序就是 default Markdown Template 的閱讀順序——使用者對著 Deliverable 逐段核對。
UNFILLED = [
    "meta.location",
    "meta.absentees",
    "topics[1].discussion",
    "topics[1].resolutions",
    "action_items[1].owner",
    "action_items[1].due",
    "action_items[1].source",
]


def run_mm(*args):
    return subprocess.run(
        [sys.executable, str(MM), *args],
        capture_output=True,
        text=True,
        encoding="utf-8",
    )


def init(root):
    result = run_mm("init", "--root", str(root))
    assert result.returncode == 0, result.stderr


def write_record(root, content=RECORD, meeting=MEETING):
    target = root / "records" / f"{meeting}.yaml"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return target


def write_template(root, content, name=TEMPLATE):
    target = root / "templates" / "markdown" / name
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return target


def write_docx_template(root, paragraphs, name="other.docx"):
    from docx import Document

    target = root / "templates" / "docx" / name
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(str(target))
    return target


def deliverable(root, meeting=MEETING):
    return root / "output" / meeting / "minutes.md"


def docx_deliverable(root, meeting=MEETING):
    return root / "output" / meeting / "minutes.docx"


def markdown_content(text):
    """把 markdown Deliverable 拆成純內容序列，拿掉標題、清單與表格的記號。

    留下的是使用者眼睛真正讀到的字，用來跟 .docx 的內容逐項對帳。
    """
    content = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if all(set(cell) == {"-"} for cell in cells):  # 表格的分隔列不是內容
                continue
            content.extend(cell for cell in cells if cell)
            continue
        content.append(line.lstrip("#").lstrip("- ").strip())
    return content


def docx_content(path):
    """把 .docx Deliverable 拆成同一種內容序列：段落與表格儲存格，照文件順序。"""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(str(path))
    content = []
    for element in document.element.body:
        if element.tag.endswith("}p"):
            text = Paragraph(element, document).text.strip()
            if text:
                content.append(text)
        elif element.tag.endswith("}tbl"):
            for row in Table(element, document).rows:
                content.extend(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
    return content


def render(root, meeting=MEETING, template=TEMPLATE, docx_template=None):
    args = ["render", meeting, "--markdown-template", template, "--root", str(root)]
    if docx_template:
        args += ["--docx-template", docx_template]
    result = run_mm(*args)
    assert result.returncode == 0, result.stderr
    return json.loads(result.stdout)


def prepared(root):
    """已跑過 init、放好 Minutes Record 的骨架。"""
    init(root)
    write_record(root)
    return root


def test_nested_and_list_fields_expand(tmp_path):
    render(prepared(tmp_path))

    markdown = deliverable(tmp_path).read_text(encoding="utf-8")

    assert "# 專案週會" in markdown
    assert "| 出席者 | 王小明、李小華、張大同 |" in markdown
    # 決議掛在自己的議題底下，不是全部倒在一起
    review, risk = markdown.split("### 2. 風險盤點")
    assert "登入頁下週上線" in review
    assert "舊版並行一週" in review
    assert "登入頁下週上線" not in risk
    # 清單型欄位一筆一列
    assert "| 補上登入頁的錯誤訊息 | 張大同 | 2026-08-04 |" in markdown
    assert "| 確認驗收時程 |" in markdown
    assert "notes/2026-07-28-project-weekly/slides.pptx.md#L12" in markdown


def test_blank_fields_render_as_unmentioned(tmp_path):
    render(prepared(tmp_path))

    markdown = deliverable(tmp_path).read_text(encoding="utf-8")

    assert "| 地點 | 未提及 |" in markdown
    assert "| 缺席者 | 未提及 |" in markdown
    # 沒有決議的議題，決議欄位仍要留下痕跡，不是整段消失
    risk = markdown.split("### 2. 風險盤點")[1]
    assert "未提及" in risk
    assert "| 確認驗收時程 | 未提及 | 未提及 | 未提及 |" in markdown


def test_unfilled_variables_are_reported(tmp_path):
    payload = render(prepared(tmp_path))

    assert payload["unfilled"] == UNFILLED


def test_template_variable_absent_from_the_record_is_unfilled(tmp_path):
    # 模板與 Schema 不綁定：模板問了 Minutes Record 沒有的欄位，也算沒填到
    init(tmp_path)
    write_record(tmp_path)
    write_template(tmp_path, "預算：{{ meta.budget or '未提及' }}\n", "budget.md.j2")

    payload = render(tmp_path, template="budget.md.j2")

    assert payload["unfilled"] == ["meta.budget"]
    assert deliverable(tmp_path).read_text(encoding="utf-8") == "預算：未提及\n"


def test_values_render_exactly_as_written(tmp_path):
    # Render 不解讀內容：日期與時間照 Minutes Record 上的寫法出現，
    # 不會被 YAML 的時間戳規則改寫成另一種格式
    init(tmp_path)
    write_record(tmp_path, "meta:\n  datetime: 2026-07-28T14:00:00+08:00\n")
    write_template(tmp_path, "{{ meta.datetime }}\n", "when.md.j2")

    render(tmp_path, template="when.md.j2")

    markdown = deliverable(tmp_path).read_text(encoding="utf-8")
    assert markdown == "2026-07-28T14:00:00+08:00\n"


def test_missing_block_is_reported_once(tmp_path):
    # 整個 meta 區塊都不在，不必把底下每個欄位各報一次
    init(tmp_path)
    write_record(tmp_path, "topics: []\n")
    write_template(tmp_path, "{{ meta.title }}/{{ meta.chair }}\n", "meta.md.j2")

    payload = render(tmp_path, template="meta.md.j2")

    assert payload["unfilled"] == ["meta"]


def test_deliverable_rebuilds_identically_after_deleting_output(tmp_path):
    prepared(tmp_path)
    first = render(tmp_path)
    content = deliverable(tmp_path).read_bytes()

    shutil.rmtree(tmp_path / "output" / MEETING)
    second = render(tmp_path)

    assert deliverable(tmp_path).read_bytes() == content
    assert second["unfilled"] == first["unfilled"]


def test_render_only_writes_the_deliverable(tmp_path):
    prepared(tmp_path)
    record = tmp_path / "records" / f"{MEETING}.yaml"
    before = record.read_bytes()
    template = (tmp_path / "templates" / "markdown" / TEMPLATE).read_bytes()

    payload = render(tmp_path)

    assert record.read_bytes() == before
    assert (tmp_path / "templates" / "markdown" / TEMPLATE).read_bytes() == template
    # Minutes Record 留在 records/，Deliverable 只在 output/
    assert list((tmp_path / "output" / MEETING).iterdir()) == [deliverable(tmp_path)]
    assert payload["deliverables"] == [str(deliverable(tmp_path))]


def test_missing_minutes_record_is_an_error(tmp_path):
    init(tmp_path)

    result = run_mm(
        "render", MEETING, "--markdown-template", TEMPLATE, "--root", str(tmp_path)
    )

    assert result.returncode == 1
    assert MEETING in result.stderr
    assert not (tmp_path / "output" / MEETING).exists()


def test_markdown_and_docx_deliverables_carry_the_same_content(tmp_path):
    # 換一份模板重新產出，兩份 Deliverable 要能互相對帳：
    # 呈現方式不同，讀到的字必須一模一樣。
    prepared(tmp_path)

    render(tmp_path, docx_template=DOCX_TEMPLATE)

    markdown = deliverable(tmp_path).read_text(encoding="utf-8")
    assert docx_content(docx_deliverable(tmp_path)) == markdown_content(markdown)


def test_nested_and_list_fields_expand_in_docx(tmp_path):
    render(prepared(tmp_path), docx_template=DOCX_TEMPLATE)

    content = docx_content(docx_deliverable(tmp_path))

    # 決議掛在自己的議題底下，不是全部倒在一起
    review = content.index("1. 進度回顧")
    risk = content.index("2. 風險盤點")
    assert any("登入頁下週上線" in text for text in content[review:risk])
    assert not any("登入頁下週上線" in text for text in content[risk:])
    # 清單型欄位一筆一列
    assert "補上登入頁的錯誤訊息" in content
    assert "確認驗收時程" in content
    assert "notes/2026-07-28-project-weekly/slides.pptx.md#L12" in " ".join(content)


def test_blank_fields_render_as_unmentioned_in_docx(tmp_path):
    render(prepared(tmp_path), docx_template=DOCX_TEMPLATE)

    content = docx_content(docx_deliverable(tmp_path))

    # 空的地點與缺席者、沒有決議的議題、沒填的待辦欄位，全都留下「未提及」
    assert content[content.index("地點") + 1] == "未提及"
    assert content[content.index("缺席者") + 1] == "未提及"
    assert content.count("未提及") == 7


def test_unfilled_covers_docx_template_variables(tmp_path):
    # 未填變數的回報涵蓋兩份模板，不只 Markdown Template
    prepared(tmp_path)
    write_docx_template(tmp_path, ["預算：{{ meta.budget or '未提及' }}"])

    payload = render(tmp_path, docx_template="other.docx")

    assert payload["unfilled"] == UNFILLED + ["meta.budget"]
    assert docx_content(docx_deliverable(tmp_path)) == ["預算：未提及"]


def test_docx_template_is_optional(tmp_path):
    # 不指定 Docx Template 就只出 markdown，這不是錯誤
    prepared(tmp_path)

    payload = render(tmp_path)

    assert payload["docx_template"] is None
    assert payload["deliverables"] == [str(deliverable(tmp_path))]
    assert not docx_deliverable(tmp_path).exists()


def test_switching_docx_template_leaves_the_minutes_record_untouched(tmp_path):
    prepared(tmp_path)
    record = tmp_path / "records" / f"{MEETING}.yaml"
    before = record.read_bytes()
    render(tmp_path, docx_template=DOCX_TEMPLATE)
    write_docx_template(tmp_path, ["{{ meta.title }}"])

    payload = render(tmp_path, docx_template="other.docx")

    assert record.read_bytes() == before
    assert docx_content(docx_deliverable(tmp_path)) == ["專案週會"]
    assert payload["docx_template"] == "other.docx"
    assert payload["deliverables"] == [
        str(deliverable(tmp_path)),
        str(docx_deliverable(tmp_path)),
    ]


def test_xml_special_characters_survive_the_docx(tmp_path):
    # .docx 內容是 XML，值裡的 & 與 < 要被跳脫，否則整份檔案打不開
    init(tmp_path)
    write_record(tmp_path, "meta:\n  title: R&D <內部> 週會\n")
    write_docx_template(tmp_path, ["{{ meta.title }}"])

    render(tmp_path, docx_template="other.docx")

    assert docx_content(docx_deliverable(tmp_path)) == ["R&D <內部> 週會"]


def test_missing_docx_template_is_an_error(tmp_path):
    prepared(tmp_path)

    result = run_mm(
        "render",
        MEETING,
        "--markdown-template",
        TEMPLATE,
        "--docx-template",
        "nope.docx",
        "--root",
        str(tmp_path),
    )

    assert result.returncode == 1
    assert "nope.docx" in result.stderr
    # 一份都不寫：markdown 不該因為 .docx 缺席而先落地
    assert not (tmp_path / "output" / MEETING).exists()


def test_missing_markdown_template_is_an_error(tmp_path):
    prepared(tmp_path)

    result = run_mm(
        "render", MEETING, "--markdown-template", "nope.md.j2", "--root", str(tmp_path)
    )

    assert result.returncode == 1
    assert "nope.md.j2" in result.stderr
    assert not (tmp_path / "output" / MEETING).exists()
