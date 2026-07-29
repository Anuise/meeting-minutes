"""`mm check` — 交付前的清單：空欄位、缺 source、模板變數對不到 Minutes Schema。

沿用同一形狀：準備 fixture 目錄、執行一次子指令、斷言 stdout 的 JSON。
"""

import json
import subprocess
import sys
from pathlib import Path

MM = Path(__file__).resolve().parent.parent / "scripts" / "mm.py"

MEETING = "2026-07-28-project-weekly"
SCHEMA = "default.yaml"
TEMPLATE = "default.md.j2"
DOCX_TEMPLATE = "default.docx"

# 有空欄位、也有一筆待辦缺 source 的 Minutes Record。正例用。
INCOMPLETE = """\
meta:
  title: 專案週會
  datetime: 2026-07-28 14:00
  location: ''
  chair: 王小明
  recorder: 李小華
  attendees: [王小明, 李小華, 張大同]
  absentees: []

topics:
  - title: 進度回顧
    discussion: 前端已完成登入頁
    resolutions:
      - text: 登入頁下週上線
        source: notes/2026-07-28-project-weekly/slides.pptx.md#L12
  - title: 風險盤點
    discussion: ''
    resolutions: []

action_items:
  - task: 補上登入頁的錯誤訊息
    owner: 張大同
    due: 2026-08-04
    source: notes/2026-07-28-project-weekly/agenda.pdf.md#L3
  - task: 確認驗收時程
    owner: '   '
    due: ''
    source:

next_meeting: 2026-08-04 14:00
"""

# default Minutes Schema 的每一個欄位都填了、每筆決議與待辦都帶 source。反例用。
COMPLETE = """\
meta:
  title: 專案週會
  datetime: 2026-07-28 14:00
  location: 三樓會議室
  chair: 王小明
  recorder: 李小華
  attendees: [王小明, 李小華, 張大同]
  absentees: [陳小美]

topics:
  - title: 進度回顧
    discussion: 前端已完成登入頁
    resolutions:
      - text: 登入頁下週上線
        source: notes/2026-07-28-project-weekly/slides.pptx.md#L12

action_items:
  - task: 補上登入頁的錯誤訊息
    owner: 張大同
    due: 2026-08-04
    source: notes/2026-07-28-project-weekly/agenda.pdf.md#L3

next_meeting: 2026-08-04 14:00
"""

# 順序就是 Minutes Schema 上的欄位順序——使用者對著 schema 逐項補。
BLANK_FIELDS = [
    "meta.location",
    "meta.absentees",
    "topics[1].discussion",
    "topics[1].resolutions",
    "action_items[1].owner",
    "action_items[1].due",
]


def run_mm(*args):
    return subprocess.run(
        [sys.executable, str(MM), *args],
        capture_output=True,
        text=True,
        encoding="utf-8",
    )


def init(root):
    result = run_mm("init", "--root", str(root))
    assert result.returncode == 0, result.stderr


def write_record(root, content, meeting=MEETING):
    target = root / "records" / f"{meeting}.yaml"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return target


def write_markdown_template(root, content, name):
    target = root / "templates" / "markdown" / name
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return target


def write_docx_template(root, paragraphs, name, header=None, footer=None):
    from docx import Document

    target = root / "templates" / "docx" / name
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    if header is not None:
        document.sections[0].header.paragraphs[0].text = header
    if footer is not None:
        document.sections[0].footer.paragraphs[0].text = footer
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(str(target))
    return target


def write_schema(root, content, name):
    target = root / "templates" / "schema" / name
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return target


def prepared(root, record=INCOMPLETE):
    """已跑過 init、放好 Minutes Record 的骨架。"""
    init(root)
    write_record(root, record)
    return root


def check_args(
    root, meeting=MEETING, schema=SCHEMA, template=TEMPLATE, docx_template=None
):
    args = [
        "check",
        meeting,
        "--schema",
        schema,
        "--markdown-template",
        template,
        "--root",
        str(root),
    ]
    if docx_template:
        args += ["--docx-template", docx_template]
    return args


def check(root, **options):
    result = run_mm(*check_args(root, **options))
    assert result.returncode == 0, result.stderr
    return json.loads(result.stdout)


def paths(findings):
    return [finding["path"] for finding in findings]


def test_blank_fields_are_listed(tmp_path):
    payload = check(prepared(tmp_path))

    assert paths(payload["blank_fields"]) == BLANK_FIELDS


def test_no_blank_fields_when_every_field_is_filled(tmp_path):
    payload = check(prepared(tmp_path, COMPLETE))

    assert payload["blank_fields"] == []


def test_missing_block_is_listed_once(tmp_path):
    # 整個 meta 區塊都不在，不必把底下每個欄位各報一次
    init(tmp_path)
    write_record(tmp_path, "topics: []\naction_items: []\nnext_meeting: 未定\n")

    payload = check(tmp_path)

    assert paths(payload["blank_fields"]) == ["meta", "topics", "action_items"]


def test_blank_field_message_names_the_field(tmp_path):
    payload = check(prepared(tmp_path))

    message = payload["blank_fields"][0]["message"]
    assert "地點" in message
    assert "meta.location" in message


def test_missing_source_is_listed(tmp_path):
    payload = check(prepared(tmp_path))

    assert paths(payload["missing_source"]) == ["action_items[1].source"]
    # 缺 source 只歸一類，不會同一件事在兩份清單上各報一次
    assert not any(path.endswith(".source") for path in paths(payload["blank_fields"]))


def test_no_missing_source_when_every_entry_cites_a_note(tmp_path):
    payload = check(prepared(tmp_path, COMPLETE))

    assert payload["missing_source"] == []


def test_missing_source_message_names_the_entry(tmp_path):
    payload = check(prepared(tmp_path))

    message = payload["missing_source"][0]["message"]
    assert "待辦事項" in message
    assert "第 2 筆" in message
    assert "action_items[1].source" in message


def test_nested_missing_source_is_listed(tmp_path):
    # 決議巢狀在議題底下，缺 source 一樣要被指名到是哪一則議題的第幾條決議
    init(tmp_path)
    write_record(
        tmp_path,
        "topics:\n"
        "  - title: 進度回顧\n"
        "    discussion: 前端已完成登入頁\n"
        "    resolutions:\n"
        "      - text: 登入頁下週上線\n"
        "        source: ''\n",
    )

    payload = check(tmp_path)

    [finding] = payload["missing_source"]
    assert finding["path"] == "topics[0].resolutions[0].source"
    assert "決議" in finding["message"]
    assert "第 1 筆" in finding["message"]


def test_markdown_template_variable_missing_from_the_schema(tmp_path):
    write_markdown_template(
        prepared(tmp_path, COMPLETE),
        "預算：{{ meta.budget or '未提及' }}\n",
        "budget.md.j2",
    )

    payload = check(tmp_path, template="budget.md.j2")

    [finding] = payload["unmapped_variables"]
    assert finding["template"] == "budget.md.j2"
    assert finding["kind"] == "markdown"
    assert finding["variable"] == "meta.budget"
    # 訊息要指名到變數、模板與 schema 三者
    assert "meta.budget" in finding["message"]
    assert "budget.md.j2" in finding["message"]
    assert SCHEMA in finding["message"]


def test_default_markdown_template_maps_onto_the_default_schema(tmp_path):
    payload = check(prepared(tmp_path, COMPLETE))

    assert payload["unmapped_variables"] == []


def test_docx_template_variable_missing_from_the_schema(tmp_path):
    write_docx_template(
        prepared(tmp_path, COMPLETE),
        ["預算：{{ meta.budget or '未提及' }}"],
        "budget.docx",
    )

    payload = check(tmp_path, docx_template="budget.docx")

    [finding] = payload["unmapped_variables"]
    assert finding["template"] == "budget.docx"
    assert finding["kind"] == "docx"
    assert finding["variable"] == "meta.budget"
    assert "meta.budget" in finding["message"]
    assert "budget.docx" in finding["message"]
    assert SCHEMA in finding["message"]


def test_default_docx_template_maps_onto_the_default_schema(tmp_path):
    payload = check(prepared(tmp_path, COMPLETE), docx_template=DOCX_TEMPLATE)

    assert payload["unmapped_variables"] == []


def test_docx_header_and_footer_variables_are_covered(tmp_path):
    # 客戶樣板常把會議名稱放在頁首、文件編號放在頁尾，那裡的變數對不到 schema 一樣是空格
    write_docx_template(
        prepared(tmp_path, COMPLETE),
        ["{{ meta.title }}"],
        "header.docx",
        header="{{ meta.client_code }}",
        footer="{{ meta.document_no }}",
    )

    payload = check(tmp_path, docx_template="header.docx")

    assert sorted(
        finding["variable"] for finding in payload["unmapped_variables"]
    ) == ["meta.client_code", "meta.document_no"]


def test_loop_variables_resolve_to_the_list_they_iterate(tmp_path):
    # 迴圈變數要換回它迭代的清單才對得上 schema，否則整份模板都會被誤報
    write_markdown_template(
        prepared(tmp_path, COMPLETE),
        "{% for topic in topics %}{{ topic.title }}{{ topic.budget }}{% endfor %}\n",
        "loop.md.j2",
    )

    payload = check(tmp_path, template="loop.md.j2")

    [finding] = payload["unmapped_variables"]
    assert finding["variable"] == "topic.budget"
    assert finding["field"] == "topics.budget"
    assert "topics.budget" in finding["message"]


def test_same_loop_variable_over_different_lists(tmp_path):
    # 兩個迴圈都把變數叫 item，迭代的卻是不同清單：後面那個對不到，不能被前面吃掉
    write_markdown_template(
        prepared(tmp_path, COMPLETE),
        "{% for item in topics %}{{ item.title }}{% endfor %}"
        "{% for item in action_items %}{{ item.title }}{% endfor %}\n",
        "twice.md.j2",
    )

    payload = check(tmp_path, template="twice.md.j2")

    [finding] = payload["unmapped_variables"]
    assert finding["variable"] == "item.title"
    assert finding["field"] == "action_items.title"


def test_indexed_expression_does_not_swallow_the_variable(tmp_path):
    # 索引運算式解不出完整欄位路徑，但底下的變數不能因此整棵被漏掉
    write_markdown_template(
        prepared(tmp_path, COMPLETE), "{{ attachments[0].name }}\n", "indexed.md.j2"
    )

    payload = check(tmp_path, template="indexed.md.j2")

    assert [finding["variable"] for finding in payload["unmapped_variables"]] == [
        "attachments"
    ]


def test_names_bound_by_set_are_not_reported(tmp_path):
    # `{% set %}` 綁出來的名字不是 Minutes Record 的欄位，報了就是誤報
    write_markdown_template(
        prepared(tmp_path, COMPLETE),
        "{% set decided = topics %}{{ decided | length }}\n",
        "set.md.j2",
    )

    payload = check(tmp_path, template="set.md.j2")

    assert payload["unmapped_variables"] == []


def test_findings_are_reported_but_do_not_block(tmp_path):
    write_markdown_template(prepared(tmp_path), "{{ meta.budget }}\n", "budget.md.j2")
    write_docx_template(tmp_path, ["{{ meta.owner_org }}"], "org.docx")

    result = run_mm(
        *check_args(tmp_path, template="budget.md.j2", docx_template="org.docx")
    )

    assert result.returncode == 0
    payload = json.loads(result.stdout)
    assert payload["blank_fields"]
    assert payload["missing_source"]
    assert len(payload["unmapped_variables"]) == 2


def test_check_writes_nothing(tmp_path):
    prepared(tmp_path)
    before = {
        path: path.read_bytes() for path in sorted(tmp_path.rglob("*")) if path.is_file()
    }

    check(tmp_path, docx_template=DOCX_TEMPLATE)

    after = {
        path: path.read_bytes() for path in sorted(tmp_path.rglob("*")) if path.is_file()
    }
    assert after == before


def test_missing_minutes_record_is_an_error(tmp_path):
    init(tmp_path)

    result = run_mm(*check_args(tmp_path))

    assert result.returncode == 1
    assert MEETING in result.stderr


def test_missing_schema_is_an_error(tmp_path):
    prepared(tmp_path)

    result = run_mm(*check_args(tmp_path, schema="nope.yaml"))

    assert result.returncode == 1
    assert "nope.yaml" in result.stderr


def test_schema_without_fields_is_an_error(tmp_path):
    # 沒有欄位的 schema 會讓檢查安靜地什麼都查不到，那比報錯更危險
    prepared(tmp_path)
    write_schema(tmp_path, "name: empty\nlabel: 空的\n", "empty.yaml")

    result = run_mm(*check_args(tmp_path, schema="empty.yaml"))

    assert result.returncode == 1
    assert "沒有定義任何欄位" in result.stderr


def test_missing_markdown_template_is_an_error(tmp_path):
    prepared(tmp_path)

    result = run_mm(*check_args(tmp_path, template="nope.md.j2"))

    assert result.returncode == 1
    assert "nope.md.j2" in result.stderr


def test_missing_docx_template_is_an_error(tmp_path):
    prepared(tmp_path)

    result = run_mm(*check_args(tmp_path, docx_template="nope.docx"))

    assert result.returncode == 1
    assert "nope.docx" in result.stderr
