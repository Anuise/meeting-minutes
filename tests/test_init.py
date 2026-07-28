"""`mm init` — 骨架與 default 模板的建立行為。

每條測試都從 CLI 接縫下手：準備一個 fixture 目錄、執行一次子指令、
斷言檔案系統的結果與 stdout 的 JSON。
"""

import json
import re
import subprocess
import sys
from pathlib import Path

import yaml

MM = Path(__file__).resolve().parent.parent / "scripts" / "mm.py"


def run_mm(*args):
    return subprocess.run(
        [sys.executable, str(MM), *args],
        capture_output=True,
        text=True,
        encoding="utf-8",
    )


def init(root):
    result = run_mm("init", "--root", str(root))
    assert result.returncode == 0, result.stderr
    return json.loads(result.stdout)


def test_version_is_reported():
    result = run_mm("--version")

    assert result.returncode == 0, result.stderr
    assert re.fullmatch(r"mm \d+\.\d+\.\d+\s*", result.stdout), result.stdout


def test_help_lists_the_planned_subcommands():
    result = run_mm("--help")

    assert result.returncode == 0, result.stderr
    for command in ("init", "ingest", "render", "check", "scan-docx", "apply-docx", "list"):
        assert command in result.stdout, f"{command} 沒有出現在 --help 中"


def test_init_creates_folder_skeleton(tmp_path):
    payload = init(tmp_path)

    expected = [
        "rawdata",
        "notes",
        "records",
        "output",
        "templates/schema",
        "templates/markdown",
        "templates/docx",
        "templates/docx-source",
    ]
    for relative in expected:
        assert (tmp_path / relative).is_dir(), f"{relative} 沒有被建立"
    assert set(expected) <= set(payload["created"])


def test_init_writes_default_minutes_schema(tmp_path):
    init(tmp_path)

    schema_file = tmp_path / "templates/schema/default.yaml"
    assert schema_file.is_file()
    schema = yaml.safe_load(schema_file.read_text(encoding="utf-8"))

    meta_keys = [field["key"] for field in schema["meta"]]
    assert meta_keys == [
        "title",
        "datetime",
        "location",
        "chair",
        "recorder",
        "attendees",
        "absentees",
    ]

    body = {field["key"]: field for field in schema["body"]}
    assert set(body) == {"topics", "action_items", "next_meeting"}

    topic_fields = {field["key"]: field for field in body["topics"]["fields"]}
    assert set(topic_fields) == {"title", "discussion", "resolutions"}
    # 決議巢狀在議題底下，以保留「哪條決議屬於哪個題目」的關係
    resolution_keys = {field["key"] for field in topic_fields["resolutions"]["fields"]}
    assert resolution_keys == {"text", "source"}

    action_keys = {field["key"] for field in body["action_items"]["fields"]}
    assert action_keys == {"task", "owner", "due", "source"}


RECORD = {
    "meta": {
        "title": "第三次專案週會",
        "datetime": "2026-07-28 14:00",
        "location": "台北 3F 會議室",
        "chair": "王小明",
        "recorder": "李大華",
        "attendees": ["王小明", "李大華"],
        "absentees": ["陳小美"],
    },
    "topics": [
        {
            "title": "API 延遲",
            "discussion": "尖峰時段回應時間拉長到兩秒。",
            "resolutions": [
                {"text": "先加上快取層", "source": "notes/2026-07-28/slides.md#L12"}
            ],
        }
    ],
    "action_items": [
        {
            "task": "導入快取層",
            "owner": "李大華",
            "due": "2026-08-04",
            "source": "notes/2026-07-28/slides.md#L14",
        }
    ],
    "next_meeting": "2026-08-04 14:00",
}


def render_markdown_template(path, record):
    from jinja2 import Environment

    environment = Environment(keep_trailing_newline=True)
    return environment.from_string(path.read_text(encoding="utf-8")).render(**record)


def test_default_markdown_template_renders_a_minutes_record(tmp_path):
    init(tmp_path)

    template = tmp_path / "templates/markdown/default.md.j2"
    assert template.is_file()

    output = render_markdown_template(template, RECORD)

    for expected in (
        "第三次專案週會",
        "台北 3F 會議室",
        "陳小美",
        "API 延遲",
        "尖峰時段回應時間拉長到兩秒。",
        "先加上快取層",
        "notes/2026-07-28/slides.md#L12",
        "導入快取層",
        "2026-08-04",
        "2026-08-04 14:00",
    ):
        assert expected in output, f"{expected} 沒有出現在渲染結果中"


def test_default_markdown_template_shows_placeholder_for_empty_fields(tmp_path):
    init(tmp_path)

    output = render_markdown_template(
        tmp_path / "templates/markdown/default.md.j2",
        {"meta": {}, "topics": [], "action_items": [], "next_meeting": None},
    )

    assert "未提及" in output


def docx_text(path):
    from docx import Document

    document = Document(str(path))
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.extend(cell.text for cell in row.cells)
    return "\n".join(lines)


def test_default_docx_template_is_renderable_by_docxtpl(tmp_path):
    from docxtpl import DocxTemplate

    init(tmp_path)

    template = tmp_path / "templates/docx/default.docx"
    assert template.is_file()

    document = DocxTemplate(str(template))
    document.render(RECORD)
    rendered = tmp_path / "rendered.docx"
    document.save(str(rendered))

    text = docx_text(rendered)
    for expected in (
        "第三次專案週會",
        "台北 3F 會議室",
        "陳小美",
        "API 延遲",
        "尖峰時段回應時間拉長到兩秒。",
        "先加上快取層",
        "notes/2026-07-28/slides.md#L12",
        "導入快取層",
        "2026-08-04",
    ):
        assert expected in text, f"{expected} 沒有出現在渲染結果中"
    assert "{{" not in text and "{%" not in text


def test_init_is_idempotent_and_never_overwrites_existing_files(tmp_path):
    init(tmp_path)

    schema = tmp_path / "templates/schema/default.yaml"
    schema.write_text("# 我自己改過的 schema\n", encoding="utf-8")
    docx = tmp_path / "templates/docx/default.docx"
    docx_bytes = docx.read_bytes()

    payload = init(tmp_path)

    assert payload["created"] == []
    assert "templates/schema/default.yaml" in payload["skipped"]
    assert "templates/docx/default.docx" in payload["skipped"]
    assert schema.read_text(encoding="utf-8") == "# 我自己改過的 schema\n"
    assert docx.read_bytes() == docx_bytes
