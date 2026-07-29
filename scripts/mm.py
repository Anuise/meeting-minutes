#!/usr/bin/env python3
"""meeting-minutes 的唯一程式進入點。

吃 argv、把結構化結果以 JSON 印到 stdout、錯誤走 exit code、完全非互動。
所有選單、確認、判斷與模型呼叫都留在 skill（agent）端。
"""

import argparse
import json
import mimetypes
import re
import sys
from pathlib import Path

VERSION = "0.1.0"

DIRECTORIES = (
    "rawdata",
    "notes",
    "records",
    "output",
    "templates/schema",
    "templates/markdown",
    "templates/docx",
    "templates/docx-source",
)


DEFAULT_SCHEMA = """\
# Default Minutes Schema：正式會議記錄型。
# 決定 Minutes Record 有哪些欄位。可以複製一份再改，做成其他會議型態的 schema。

name: default
label: 正式會議記錄

meta:
  - { key: title, label: 會議名稱, type: text }
  - { key: datetime, label: 日期時間, type: datetime }
  - { key: location, label: 地點, type: text }
  - { key: chair, label: 主席, type: text }
  - { key: recorder, label: 記錄人, type: text }
  - { key: attendees, label: 出席者, type: people }
  - { key: absentees, label: 缺席者, type: people }

body:
  - key: topics
    label: 議題
    type: list
    fields:
      - { key: title, label: 題目, type: text }
      - { key: discussion, label: 討論摘要, type: text }
      # 決議巢狀在議題底下，保留「哪條決議屬於哪個題目」的關係
      - key: resolutions
        label: 決議
        type: list
        fields:
          - { key: text, label: 決議內容, type: text }
          - { key: source, label: 來源, type: source }

  - key: action_items
    label: 待辦事項
    type: list
    fields:
      - { key: task, label: 事項, type: text }
      - { key: owner, label: 負責人, type: text }
      - { key: due, label: 期限, type: date }
      - { key: source, label: 來源, type: source }

  - { key: next_meeting, label: 下次會議, type: text }
"""

DEFAULT_MARKDOWN_TEMPLATE = """\
# {{ meta.title or '未提及' }}

| 項目 | 內容 |
| --- | --- |
| 日期時間 | {{ meta.datetime or '未提及' }} |
| 地點 | {{ meta.location or '未提及' }} |
| 主席 | {{ meta.chair or '未提及' }} |
| 記錄人 | {{ meta.recorder or '未提及' }} |
| 出席者 | {{ meta.attendees | default([], true) | join('、') or '未提及' }} |
| 缺席者 | {{ meta.absentees | default([], true) | join('、') or '未提及' }} |

## 議題
{% for topic in topics | default([], true) %}
### {{ loop.index }}. {{ topic.title or '未提及' }}

{{ topic.discussion or '未提及' }}

決議：
{%- for resolution in topic.resolutions | default([], true) %}
- {{ resolution.text or '未提及' }}（來源：{{ resolution.source or '未提及' }}）
{%- else %}
- 未提及
{%- endfor %}
{% else %}
未提及
{% endfor %}
## 待辦事項

| 事項 | 負責人 | 期限 | 來源 |
| --- | --- | --- | --- |
{%- for item in action_items | default([], true) %}
| {{ item.task or '未提及' }} | {{ item.owner or '未提及' }} | \
{{ item.due or '未提及' }} | {{ item.source or '未提及' }} |
{%- else %}
| 未提及 | 未提及 | 未提及 | 未提及 |
{%- endfor %}

## 下次會議

{{ next_meeting or '未提及' }}
"""

DOCX_META_ROWS = (
    ("日期時間", "{{ meta.datetime or '未提及' }}"),
    ("地點", "{{ meta.location or '未提及' }}"),
    ("主席", "{{ meta.chair or '未提及' }}"),
    ("記錄人", "{{ meta.recorder or '未提及' }}"),
    ("出席者", "{{ meta.attendees | default([], true) | join('、') or '未提及' }}"),
    ("缺席者", "{{ meta.absentees | default([], true) | join('、') or '未提及' }}"),
)

DOCX_ACTION_COLUMNS = (
    ("事項", "{{ item.task or '未提及' }}"),
    ("負責人", "{{ item.owner or '未提及' }}"),
    ("期限", "{{ item.due or '未提及' }}"),
    ("來源", "{{ item.source or '未提及' }}"),
)


def write_default_docx(target):
    """產生一份已標上 Jinja2 變數、能被 docxtpl 直接渲染的 default Docx Template。

    內容逐項對齊 default Markdown Template：同一份 Minutes Record 套這兩份模板，
    兩份 Deliverable 上讀到的字要一模一樣，使用者才對得了帳。
    """
    from docx import Document

    document = Document()
    document.add_paragraph("{{ meta.title or '未提及' }}", style="Title")

    meta_table = document.add_table(rows=1, cols=2)
    meta_table.style = "Table Grid"
    meta_header = meta_table.rows[0].cells
    meta_header[0].text = "項目"
    meta_header[1].text = "內容"
    for label, expression in DOCX_META_ROWS:
        cells = meta_table.add_row().cells
        cells[0].text = label
        cells[1].text = expression

    document.add_heading("議題", level=1)
    document.add_paragraph("{%p for topic in topics | default([], true) %}")
    document.add_heading("{{ loop.index }}. {{ topic.title or '未提及' }}", level=2)
    document.add_paragraph("{{ topic.discussion or '未提及' }}")
    document.add_paragraph("決議：")
    document.add_paragraph(
        "{%p for resolution in topic.resolutions | default([], true) %}"
    )
    document.add_paragraph(
        "{{ resolution.text or '未提及' }}（來源：{{ resolution.source or '未提及' }}）",
        style="List Bullet",
    )
    # 沒有決議的議題也要留下痕跡，跟 markdown 一樣落一行「未提及」
    document.add_paragraph("{%p else %}")
    document.add_paragraph("未提及", style="List Bullet")
    document.add_paragraph("{%p endfor %}")
    document.add_paragraph("{%p else %}")
    document.add_paragraph("未提及")
    document.add_paragraph("{%p endfor %}")

    document.add_heading("待辦事項", level=1)
    action_table = document.add_table(rows=1, cols=len(DOCX_ACTION_COLUMNS))
    action_table.style = "Table Grid"
    header = action_table.rows[0].cells
    for index, (label, _) in enumerate(DOCX_ACTION_COLUMNS):
        header[index].text = label
    action_table.add_row().cells[0].text = (
        "{%tr for item in action_items | default([], true) %}"
    )
    repeated_row = action_table.add_row().cells
    for index, (_, expression) in enumerate(DOCX_ACTION_COLUMNS):
        repeated_row[index].text = expression
    action_table.add_row().cells[0].text = "{%tr else %}"
    for cell in action_table.add_row().cells:
        cell.text = "未提及"
    action_table.add_row().cells[0].text = "{%tr endfor %}"

    document.add_heading("下次會議", level=1)
    document.add_paragraph("{{ next_meeting or '未提及' }}")

    document.save(str(target))


def write_text(content):
    return lambda target: target.write_text(content, encoding="utf-8")


# 每個 default 檔案配一個 writer。docx 得用程式產生，其餘是固定文字。
SEEDS = (
    ("templates/schema/default.yaml", write_text(DEFAULT_SCHEMA)),
    ("templates/markdown/default.md.j2", write_text(DEFAULT_MARKDOWN_TEMPLATE)),
    ("templates/docx/default.docx", write_default_docx),
)


def cmd_init(args):
    """建立骨架與 default 模板。已存在的一律跳過，絕不覆蓋使用者改過的內容。"""
    root = Path(args.root)
    created = []
    skipped = []

    for relative in DIRECTORIES:
        target = root / relative
        if target.is_dir():
            skipped.append(relative)
        else:
            target.mkdir(parents=True)
            created.append(relative)

    for relative, write in SEEDS:
        target = root / relative
        if target.exists():
            skipped.append(relative)
        else:
            write(target)
            created.append(relative)

    return {"root": str(root), "created": created, "skipped": skipped}


class CommandError(Exception):
    """使用者能修正的錯誤：訊息走 stderr、走非零 exit code。"""

    EXIT_CODE = 1  # 2 是 argparse 的用法錯誤，留給它


# ADR-0003：不收音訊、影片。ADR-0005：不收圖片。
# 兩份清單是地板——寫死才不會隨基底映像的 mime 表版本漂移；mimetypes 只做補漏。
AUDIO_SUFFIXES = frozenset(
    {
        ".mp3", ".m4a", ".m4b", ".wav", ".flac", ".aac", ".ogg", ".opus",
        ".wma", ".aiff", ".amr",
        ".mp4", ".m4v", ".mov", ".mkv", ".avi", ".webm", ".wmv", ".flv",
    }
)
IMAGE_SUFFIXES = frozenset(
    {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tif", ".tiff", ".heic"}
)

AUDIO_MESSAGE = "不支援音訊與影片。請先自行轉成逐字稿，再把逐字稿放進 rawdata/。"
IMAGE_MESSAGE = (
    "不轉圖片：容器內沒有文字辨識能力，轉出來只會是空的 Note。"
    "照片上的內容需要進會議記錄的話，請自行補一份文字說明放進 rawdata/。"
)


def unsupported_message(raw):
    """回傳「這個檔案為什麼不轉」；該轉的回 None。"""
    suffix = raw.suffix.lower()
    if suffix in AUDIO_SUFFIXES:
        return AUDIO_MESSAGE
    if suffix in IMAGE_SUFFIXES:
        return IMAGE_MESSAGE

    mimetype = mimetypes.guess_type(raw.name)[0] or ""
    if mimetype.startswith(("audio/", "video/")):
        return AUDIO_MESSAGE
    if mimetype.startswith("image/"):
        return IMAGE_MESSAGE
    return None


def cmd_ingest(args):
    """把 rawdata/<meeting>/ 底下的檔案機械轉成 notes/<meeting>/ 的 markdown。

    只寫 notes/，絕不寫 rawdata/。單一檔案失敗不中斷整批，失敗項目集中回報。
    """
    from markitdown import MarkItDown

    root = Path(args.root)
    raw_root = root / "rawdata" / args.meeting
    if not raw_root.is_dir():
        raise CommandError(f"找不到 Raw Material 目錄：rawdata/{args.meeting}")
    note_root = root / "notes" / args.meeting

    markitdown = MarkItDown()
    ingested = []
    skipped = []
    unsupported = []
    failed = []

    for raw in sorted(path for path in raw_root.rglob("*") if path.is_file()):
        relative = raw.relative_to(raw_root).as_posix()

        message = unsupported_message(raw)
        if message:
            unsupported.append({"raw": relative, "message": message})
            continue

        # 副檔名留在 Note 檔名裡：看得出來源，且 slides.pdf 與 slides.docx 不會互相蓋掉
        note_relative = f"{relative}.md"
        note = note_root / note_relative
        # 只在 Note 確實比 Raw Material 新時跳過：時間戳打平時寧可多轉一次，
        # 也不要漏掉使用者剛換上去的素材。
        if note.exists() and note.stat().st_mtime > raw.stat().st_mtime:
            skipped.append({"raw": relative, "note": note_relative})
            continue

        try:
            markdown = markitdown.convert(str(raw)).markdown
        except Exception as error:
            failed.append({"raw": relative, "error": f"{type(error).__name__}: {error}"})
            continue

        note.parent.mkdir(parents=True, exist_ok=True)
        note.write_text(markdown, encoding="utf-8")
        ingested.append({"raw": relative, "note": note_relative})

    return {
        "meeting": args.meeting,
        "rawdata": str(raw_root),
        "notes": str(note_root),
        "ingested": ingested,
        "skipped": skipped,
        "unsupported": unsupported,
        "failed": failed,
    }


def has_files(directory):
    """這一階段真的做過嗎。空資料夾不算——使用者建了資料夾但還沒放東西。"""
    return directory.is_dir() and any(path.is_file() for path in directory.rglob("*"))


def filenames(directory, pattern):
    if not directory.is_dir():
        return []
    return sorted(
        path.name
        for path in directory.glob(pattern)
        # 使用者在 Word 裡開著模板時會留下 ~$ 開頭的鎖檔，那不是模板
        if path.is_file() and not path.name.startswith("~$")
    )


def cmd_list(args):
    """回報每個 Meeting 走到哪一步，以及目前有哪些 schema 與模板可用。

    只讀不寫。骨架還沒建起來也不算錯——回空清單就好。
    """
    root = Path(args.root)

    # 四個階段目錄各自獨立：只要任何一個有這個 slug，這場 Meeting 就得被列出來
    slugs = set()
    for area in ("rawdata", "notes", "output"):
        area_root = root / area
        if area_root.is_dir():
            slugs.update(path.name for path in area_root.iterdir() if path.is_dir())
    records_root = root / "records"
    if records_root.is_dir():
        slugs.update(path.stem for path in records_root.glob("*.yaml") if path.is_file())

    meetings = [
        {
            "slug": slug,
            "raw_material": has_files(root / "rawdata" / slug),
            "note": has_files(root / "notes" / slug),
            "minutes_record": (root / "records" / f"{slug}.yaml").is_file(),
            "deliverable": has_files(root / "output" / slug),
        }
        for slug in sorted(slugs)
    ]

    return {
        "root": str(root),
        "meetings": meetings,
        # 三個清單都按副檔名過濾，跟 init 產出的 default 對齊：
        # README、.gitkeep、Word 的鎖檔不該被當成可用模板交給 render
        "schemas": filenames(root / "templates" / "schema", "*.yaml"),
        "markdown_templates": filenames(root / "templates" / "markdown", "*.j2"),
        # docx-source 底下的還沒打洞，不能拿來渲染，所以不算可用模板
        "docx_templates": filenames(root / "templates" / "docx", "*.docx"),
    }


def blank(value):
    """這個欄位沒被填到嗎。Extract 抓不到的欄位一律留空，各種「空」在這裡收斂成一個判斷。"""
    if isinstance(value, str):
        return not value.strip()
    if isinstance(value, (list, dict)):
        return not value
    return value is None


def build_environment(data, unfilled, **options):
    """建出會記下未填變數的 Jinja2 環境，以及它要吃的 context。

    未填變數以「模板真的讀到」為準——那才是 Deliverable 上看得到的空格。
    Minutes Record 裡有但模板沒用到的欄位不算，模板問了而 Minutes Record 沒有的算。
    兩種模板共用這一套，所以 unfilled 涵蓋 Markdown 與 Docx Template 兩邊的變數。
    """
    from jinja2 import ChainableUndefined, Environment

    class Unfilled(ChainableUndefined):
        """沒填到的變數。渲染成空字串（模板的 `or '未提及'` 接手），並記下自己的路徑。"""

        def __init__(self, *args, name=None, **kwargs):
            super().__init__(*args, name=name, **kwargs)
            if name and name not in unfilled:
                unfilled.append(name)

    class Block(dict):
        """Minutes Record 的一個區塊。子欄位在被模板讀到的那一刻才判斷有沒有填。"""

        def __init__(self, fields, path):
            super().__init__(fields)
            self._path = path

        def __getitem__(self, key):
            # 缺 key 與填了空值一視同仁：模板問了，Deliverable 上就是一個空格
            return wrap(self.get(key), f"{self._path}.{key}")

    def wrap(value, path):
        if blank(value):
            return Unfilled(name=path)
        if isinstance(value, dict):
            return Block(value, path)
        if isinstance(value, list):
            return [wrap(item, f"{path}[{index}]") for index, item in enumerate(value)]
        return value

    environment = Environment(undefined=Unfilled, **options)
    # 空的最上層欄位不放進 context，留給 Unfilled 去接：
    # 這樣它只在模板真的讀到時才被記下來，順序也跟著模板的閱讀順序。
    context = {key: wrap(value, key) for key, value in data.items() if not blank(value)}
    return environment, context


def load_record(root, meeting, action):
    """讀 Minutes Record，回傳 (路徑, 內容)。

    Minutes Record 允許人工編修，所以缺檔與壞掉的 YAML 都是使用者修得好的錯誤，
    不是 bug。BaseLoader 讓所有純量都留在字串：日期與時間照 Minutes Record 上的
    寫法處理，不會被 YAML 的時間戳規則改寫成另一種格式。
    """
    import yaml

    record = root / "records" / f"{meeting}.yaml"
    if not record.is_file():
        raise CommandError(
            f"找不到 Minutes Record：records/{meeting}.yaml。"
            f"先做 Extract 產出它，再來 {action}。"
        )
    try:
        data = yaml.load(record.read_text(encoding="utf-8"), yaml.BaseLoader) or {}
    except yaml.YAMLError as error:
        raise CommandError(f"Minutes Record 不是合法的 YAML：{error}") from error
    if not isinstance(data, dict):
        raise CommandError("Minutes Record 的最上層必須是欄位，例如 meta:、topics:。")
    return record, data


def render_markdown(template_text, data, unfilled):
    """把 Minutes Record 套上 Markdown Template，回傳 markdown 內容。"""
    environment, context = build_environment(data, unfilled, keep_trailing_newline=True)
    return environment.from_string(template_text).render(context)


def render_docx(template, target, data, unfilled):
    """把 Minutes Record 套上 Docx Template，寫出 .docx。

    autoescape 是必要的：.docx 的內容是 XML，值裡的 & 與 < 沒跳脫就整份打不開。
    """
    from docxtpl import DocxTemplate

    environment, context = build_environment(data, unfilled, autoescape=True)
    document = DocxTemplate(str(template))
    document.render(context, environment)
    document.save(str(target))


def cmd_render(args):
    """把 Minutes Record 套上 Markdown Template 與（選填的）Docx Template，寫出 Deliverable。

    不呼叫模型，只讀 records/ 與 templates/、只寫 output/，跑幾次結果都一樣。
    """
    root = Path(args.root)
    record, data = load_record(root, args.meeting, "Render")
    template = root / "templates" / "markdown" / args.markdown_template
    if not template.is_file():
        raise CommandError(
            f"找不到 Markdown Template：templates/markdown/{args.markdown_template}"
        )
    # 沒指定 Docx Template 就只出 markdown，那不是錯誤。
    # 兩份模板都在寫檔前先查，缺一份就一份都不落地。
    docx_template = None
    if args.docx_template:
        docx_template = root / "templates" / "docx" / args.docx_template
        if not docx_template.is_file():
            raise CommandError(
                f"找不到 Docx Template：templates/docx/{args.docx_template}"
            )

    unfilled = []
    markdown = render_markdown(template.read_text(encoding="utf-8"), data, unfilled)

    deliverable = root / "output" / args.meeting / "minutes.md"
    deliverable.parent.mkdir(parents=True, exist_ok=True)
    deliverable.write_text(markdown, encoding="utf-8")
    deliverables = [deliverable]

    if docx_template:
        docx_deliverable = deliverable.parent / "minutes.docx"
        render_docx(docx_template, docx_deliverable, data, unfilled)
        deliverables.append(docx_deliverable)

    return {
        "meeting": args.meeting,
        "minutes_record": str(record),
        "markdown_template": args.markdown_template,
        "docx_template": args.docx_template,
        "deliverables": [str(path) for path in deliverables],
        "unfilled": unfilled,
    }


def schema_fields(entries):
    """把 Minutes Schema 的一段欄位定義攤成 key -> 節點。

    節點帶 label（訊息要指名到欄位）、type（`source` 型別另外歸類）與子欄位。
    """
    fields = {}
    for entry in entries or []:
        key = entry.get("key")
        if key:
            fields[key] = {
                "label": entry.get("label", key),
                "type": entry.get("type", "text"),
                "fields": schema_fields(entry.get("fields")),
            }
    return fields


def load_schema(path):
    """讀 Minutes Schema，回傳與 Minutes Record 同形狀的欄位樹。

    `meta` 在 Minutes Record 裡收在 `meta.<key>` 底下，`body` 底下的區塊則攤平在
    最上層——這裡照同一個形狀組，路徑才對得上 Minutes Record 與模板變數。
    """
    import yaml

    try:
        document = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    except yaml.YAMLError as error:
        raise CommandError(f"Minutes Schema 不是合法的 YAML：{error}") from error
    if not isinstance(document, dict):
        raise CommandError("Minutes Schema 的最上層必須是欄位，例如 meta:、body:。")

    fields = {}
    meta = schema_fields(document.get("meta"))
    if meta:
        fields["meta"] = {"label": "會議資訊", "type": "block", "fields": meta}
    fields.update(schema_fields(document.get("body")))
    if not fields:
        raise CommandError("Minutes Schema 沒有定義任何欄位，meta: 與 body: 都是空的。")
    return fields


def scan_record(fields, data, path, owner, blank_fields, missing_source):
    """依 Minutes Schema 逐個欄位看 Minutes Record 填了沒。

    以 Schema 為準而不是以 Minutes Record 為準：Minutes Record 裡根本沒有的欄位
    才是最容易被漏掉的那一種。`owner` 是所在清單項目的說法（例如「決議」第 2 筆），
    缺 source 的訊息要靠它指名到是哪一筆。
    """
    for key, node in fields.items():
        value = data.get(key) if isinstance(data, dict) else None
        here = f"{path}.{key}" if path else key

        if blank(value):
            if node["type"] == "source":
                # 缺 source 比一般空欄位嚴重：查不回 Note，別人質疑時無從對證
                where = owner or f"「{node['label']}」"
                missing_source.append(
                    {
                        "path": here,
                        "label": node["label"],
                        "message": (
                            f"Minutes Record 的{where}沒有 source，指不回 Note：{here}"
                        ),
                    }
                )
            else:
                blank_fields.append(
                    {
                        "path": here,
                        "label": node["label"],
                        "message": f"Minutes Record 的「{node['label']}」是空的：{here}",
                    }
                )
            continue

        if node["type"] == "list":
            for index, item in enumerate(value if isinstance(value, list) else []):
                scan_record(
                    node["fields"],
                    item,
                    f"{here}[{index}]",
                    f"「{node['label']}」第 {index + 1} 筆",
                    blank_fields,
                    missing_source,
                )
        elif node["fields"]:
            scan_record(node["fields"], value, here, owner, blank_fields, missing_source)


# Jinja2 自己提供的名字，不是 Minutes Record 的欄位，別拿去跟 Minutes Schema 對帳
JINJA_GLOBALS = frozenset(
    {"loop", "range", "dict", "lipsum", "cycler", "joiner", "namespace"}
)


def resolve_path(node, scope):
    """把一段運算式化成 (模板上的寫法, Minutes Schema 上的欄位路徑)。

    迴圈變數換回它迭代的清單，所以 `resolution.text` 會解到
    `topics.resolutions.text`——那才是拿去跟 Minutes Schema 對帳的東西。
    解不出來的（例如 `{% set %}` 綁出來的名字）欄位路徑回 None，代表不該報。
    """
    from jinja2 import nodes

    if isinstance(node, nodes.Name):
        return node.name, scope[node.name] if node.name in scope else node.name
    if isinstance(node, nodes.Getattr):
        written, field = resolve_path(node.node, scope)
        if written is None:
            return None, None
        return f"{written}.{node.attr}", f"{field}.{node.attr}" if field else None
    if isinstance(node, nodes.Filter):
        # `topic.resolutions | default([], true)` 問的還是 topic.resolutions
        return resolve_path(node.node, scope) if node.node else (None, None)
    return None, None


def target_names(node):
    """`{% for topic in ... %}` 或 `{% set a, b = ... %}` 綁出來的名字。"""
    from jinja2 import nodes

    if isinstance(node, nodes.Name):
        return [node.name]
    return [child.name for child in node.find_all(nodes.Name)]


def add_reference(written, field, references):
    """記下模板讀到的一個欄位。

    去重要連欄位路徑一起看：兩個迴圈都把變數叫 `item`、迭代的卻是不同清單時，
    那是兩筆不同的參照，併成一筆會讓其中一份模板的問題被靜默漏掉。
    """
    if written.split(".")[0] in JINJA_GLOBALS:
        return
    reference = {"variable": written, "field": field}
    if reference not in references:
        references.append(reference)


def scan_template_node(node, scope, references):
    """走 Jinja2 的語法樹，收集模板真的讀到的欄位路徑。

    只收最深的那一層：`meta.title` 記一筆，不會連 `meta` 也記一筆。
    解不出欄位路徑的寫法（索引運算式、`{% set %}` 綁的名字）就繼續往下走——
    外層寫法特別不該讓底下的變數整棵被漏掉。
    """
    from jinja2 import nodes

    if isinstance(node, nodes.For):
        written, field = resolve_path(node.iter, scope)
        if field:
            add_reference(written, field, references)
        else:
            scan_template_node(node.iter, scope, references)
        inner = dict(scope)
        for name in target_names(node.target):
            inner[name] = field
        for child in ([node.test] if node.test else []) + node.body + node.else_:
            scan_template_node(child, inner, references)
        return

    if isinstance(node, nodes.Assign):
        scan_template_node(node.node, scope, references)
        # `{% set %}` 綁出來的名字不是 Minutes Record 的欄位，別誤報
        for name in target_names(node.target):
            scope[name] = None
        return

    if isinstance(node, (nodes.Name, nodes.Getattr)):
        written, field = resolve_path(node, scope)
        if field:
            add_reference(written, field, references)
            return

    for child in node.iter_child_nodes():
        scan_template_node(child, scope, references)


def unmapped(fields, field):
    """這條欄位路徑在 Minutes Schema 裡找不到對應嗎。"""
    for key in field.split("."):
        node = fields.get(key)
        if node is None:
            return True
        fields = node["fields"]
    return False


def docx_template_source(path):
    """把 Docx Template 攤成 docxtpl 真正拿去渲染的那串 Jinja2 原始碼。

    含頁首頁尾——客戶樣板常把會議名稱放在頁首，那裡的變數對不到 schema 一樣是空格。
    """
    from docxtpl import DocxTemplate

    template = DocxTemplate(str(path))
    template.init_docx()
    source = template.patch_xml(template.get_xml())
    for uri in (template.HEADER_URI, template.FOOTER_URI):
        for _, part in template.get_headers_footers(uri):
            source += template.patch_xml(template.get_part_xml(part))
    return source


def scan_template(source, name, kind, label, fields, schema, findings):
    """列出這份模板裡對不到 Minutes Schema 的變數。

    訊息要同時指名變數、模板與 schema：三者不強制綁定，少講一個使用者就查不下去。
    """
    from jinja2 import Environment

    references = []
    scan_template_node(Environment().parse(source), {}, references)

    for reference in references:
        variable, field = reference["variable"], reference["field"]
        if not unmapped(fields, field):
            continue
        written = variable if variable == field else f"{variable}（欄位路徑 {field}）"
        findings.append(
            {
                "template": name,
                "kind": kind,
                "variable": variable,
                "field": field,
                "message": (
                    f"{label}「{name}」用到的變數 {written}，"
                    f"在 Minutes Schema「{schema}」裡找不到對應欄位。"
                ),
            }
        )


def cmd_check(args):
    """交付前的清單：空欄位、缺 source、模板變數對不到 Minutes Schema。

    只讀不寫，只列清單不阻擋——有發現也照樣 exit 0，由使用者自己判斷。
    """
    root = Path(args.root)
    record, data = load_record(root, args.meeting, "Check")
    schema = root / "templates" / "schema" / args.schema
    if not schema.is_file():
        raise CommandError(f"找不到 Minutes Schema：templates/schema/{args.schema}")
    template = root / "templates" / "markdown" / args.markdown_template
    if not template.is_file():
        raise CommandError(
            f"找不到 Markdown Template：templates/markdown/{args.markdown_template}"
        )
    # 沒指定 Docx Template 就只查 Markdown Template，那不是錯誤。
    docx_template = None
    if args.docx_template:
        docx_template = root / "templates" / "docx" / args.docx_template
        if not docx_template.is_file():
            raise CommandError(
                f"找不到 Docx Template：templates/docx/{args.docx_template}"
            )

    fields = load_schema(schema)

    blank_fields = []
    missing_source = []
    scan_record(fields, data, "", None, blank_fields, missing_source)

    unmapped_variables = []
    scan_template(
        template.read_text(encoding="utf-8"),
        args.markdown_template,
        "markdown",
        "Markdown Template",
        fields,
        args.schema,
        unmapped_variables,
    )
    if docx_template:
        scan_template(
            docx_template_source(docx_template),
            args.docx_template,
            "docx",
            "Docx Template",
            fields,
            args.schema,
            unmapped_variables,
        )

    return {
        "meeting": args.meeting,
        "minutes_record": str(record),
        "schema": args.schema,
        "markdown_template": args.markdown_template,
        "docx_template": args.docx_template,
        "blank_fields": blank_fields,
        "missing_source": missing_source,
        "unmapped_variables": unmapped_variables,
    }


def containers(document):
    """所有可能藏著文字的容器：本文，以及有自己定義的頁首頁尾。

    首頁與偶數頁的頁首頁尾各自獨立，得逐個問。`is_linked_to_previous` 的沒有自己的
    定義（內容來自前一節），碰它只會憑空生出一份空頁首，所以跳過。
    """
    yield "body", document
    for section in document.sections:
        parts = (
            ("header", section.header),
            ("header", section.first_page_header),
            ("header", section.even_page_header),
            ("footer", section.footer),
            ("footer", section.first_page_footer),
            ("footer", section.even_page_footer),
        )
        for where, part in parts:
            if not part.is_linked_to_previous:
                yield where, part


def walk(container, where, cell=None):
    """走訪一個容器裡的所有段落，帶著它在文件上的位置。

    表格儲存格與巢狀表格都涵蓋。`cell` 是最內層儲存格的 (表格, 列, 欄)。
    """
    for paragraph in container.paragraphs:
        yield where, cell, paragraph
    for index, table in enumerate(container.tables):
        # 存元素本身而不是 id()：lxml 的 proxy 被回收後位址會被下一個元素重用，
        # 拿 id() 當鍵會把不同的儲存格誤認成同一個而漏掉。
        seen = set()
        for row_index, row in enumerate(table.rows):
            for column_index, current in enumerate(row.cells):
                # 合併過的儲存格會在同一列裡重複出現，同一個 <w:tc> 只走一次
                if current._tc in seen:
                    continue
                seen.add(current._tc)
                yield from walk(current, where, (index, row_index, column_index))


def walk_document(document):
    for where, container in containers(document):
        yield from walk(container, where)


def open_docx_source(root, name):
    from docx import Document

    source = root / "templates" / "docx-source" / name
    if not source.is_file():
        raise CommandError(f"找不到 Docx Source：templates/docx-source/{name}")
    return source, Document(str(source))


def cmd_scan_docx(args):
    """列出 Docx Source 上的每一段文字，供 agent 判斷哪些是變動欄位。

    只讀不寫。文字以段落為單位——Word 把一句話拆成幾個 run 是它的事，
    使用者眼裡那是一句話，對照表上也該是一句話。
    """
    root = Path(args.root)
    source, document = open_docx_source(root, args.source)

    items = []
    seen = set()
    for where, cell, paragraph in walk_document(document):
        text = paragraph.text.strip()
        # 對照表以文字為鍵，重複的文字列第二次只是噪音
        if not text or text in seen:
            continue
        seen.add(text)
        item = {"where": where, "kind": "cell" if cell else "paragraph", "text": text}
        if cell:
            item |= {"table": cell[0], "row": cell[1], "column": cell[2]}
        items.append(item)

    return {"source": args.source, "path": str(source), "items": items}


def read_mapping(stream):
    """從 stdin 讀最終的變數對照表。

    對照表的長度沒有上限，塞進 argv 遲早會炸，所以走 stdin——仍然完全非互動。
    """
    raw = stream.read().decode("utf-8")
    if not raw.strip():
        raise CommandError(
            "對照表是空的。apply-docx 從 stdin 讀 JSON，"
            '格式為 [{"text": "原文", "variable": "{{ 變數 }}"}]。'
        )
    try:
        entries = json.loads(raw)
    except json.JSONDecodeError as error:
        raise CommandError(f"對照表不是合法的 JSON：{error}") from error

    if not isinstance(entries, list) or not entries:
        raise CommandError("對照表必須是至少一項的 JSON 陣列。")

    mapping = {}
    for entry in entries:
        if not isinstance(entry, dict):
            raise CommandError(f"對照表的每一項都必須是物件，收到：{entry!r}")
        text = entry.get("text")
        variable = entry.get("variable")
        if not isinstance(text, str) or not text.strip():
            raise CommandError(f"對照表有一項缺少 text 或 text 是空的：{entry!r}")
        if not isinstance(variable, str) or not variable.strip():
            raise CommandError(f"對照表有一項缺少 variable：{entry!r}")
        # 同一段原文對到兩個變數，該聽誰的？寧可停下來讓使用者決定
        if text in mapping:
            raise CommandError(f"對照表裡同一段原文出現兩次：{text}")
        mapping[text] = variable

    return mapping


def punch(paragraph, pattern, mapping):
    """就地把段落裡的原文換成 Jinja2 變數，回傳這一段打掉了哪些原文。

    先在段落層級合併 run 再替換：Word 常把一句話拆成好幾個 run，
    不合併的話跨 run 的字串永遠比對不到。代價是段落內各 run 的差異
    （例如標籤粗體、值不粗體）會被第一個 run 的格式統一掉。

    段落的文字有一部分不在 run 裡（超連結、內容控制項）時整段跳過：合併只搬得動
    run，硬做會把那些文字擠到別的位置去。那些原文會出現在 unmatched 裡，
    使用者看得到自己有一段沒被打洞。
    """
    runs = paragraph.runs
    text = "".join(run.text for run in runs)
    if text != paragraph.text:
        return []

    hits = pattern.findall(text)
    if not hits:
        return []

    runs[0].text = pattern.sub(lambda match: mapping[match.group(0)], text)
    for run in runs[1:]:
        run.text = ""
    return hits


def cmd_apply_docx(args):
    """依對照表把 Docx Source 打洞成 Docx Template，輸出到 templates/docx/。

    只讀 templates/docx-source/、只寫 templates/docx/：原檔一個位元都不動，
    使用者隨時能改一改對照表重新來過。頁首頁尾、logo、字型與表格樣式都留在原地，
    因為打洞只換段落裡的文字，不重建文件。
    """
    from collections import Counter

    root = Path(args.root)
    output = args.output or args.source
    if not output.endswith(".docx"):
        raise CommandError(f"Docx Template 的檔名必須以 .docx 結尾：{output}")

    source, document = open_docx_source(root, args.source)
    mapping = read_mapping(sys.stdin.buffer)

    target = root / "templates" / "docx" / output
    if target.exists() and not args.force:
        raise CommandError(
            f"templates/docx/{output} 已經存在。"
            "換一個 --output 名稱，或加 --force 覆蓋它。"
        )

    # 長的原文先比對：不然「王小明」會先吃掉「王小明副總」的前三個字。
    # 一次 sub 掃完整段，剛換上去的 {{ 變數 }} 不會再被後面的原文比到。
    pattern = re.compile(
        "|".join(re.escape(text) for text in sorted(mapping, key=len, reverse=True))
    )

    counts = Counter()
    for _, _, paragraph in walk_document(document):
        counts.update(punch(paragraph, pattern, mapping))

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))

    return {
        "source": args.source,
        "path": str(source),
        "docx_template": str(target),
        "replaced": [
            {"text": text, "variable": variable, "count": counts[text]}
            for text, variable in mapping.items()
            if counts[text]
        ],
        # 對不到的原文要講出來：使用者以為那一格打了洞，其實原文還留在模板上
        "unmatched": [text for text in mapping if not counts[text]],
    }


def build_parser():
    parser = argparse.ArgumentParser(prog="mm", description="meeting-minutes CLI")
    parser.add_argument("--version", action="version", version=f"mm {VERSION}")
    subparsers = parser.add_subparsers(dest="command", required=True)

    init = subparsers.add_parser("init", help="建立資料夾骨架與 default 模板")
    init.add_argument("--root", default="/work", help="骨架的根目錄（預設 /work）")
    init.set_defaults(func=cmd_init)

    ingest = subparsers.add_parser("ingest", help="把 Raw Material 轉成 Note")
    ingest.add_argument("meeting", help="Meeting slug，即 rawdata/ 底下的資料夾名稱")
    ingest.add_argument("--root", default="/work", help="骨架的根目錄（預設 /work）")
    ingest.set_defaults(func=cmd_ingest)

    render = subparsers.add_parser("render", help="把 Minutes Record 套模板變成 Deliverable")
    render.add_argument("meeting", help="Meeting slug，即 records/ 底下的檔名（不含 .yaml）")
    render.add_argument(
        "--markdown-template",
        required=True,
        help="templates/markdown/ 底下的檔名，例如 default.md.j2",
    )
    render.add_argument(
        "--docx-template",
        help="templates/docx/ 底下的檔名，例如 default.docx。不給就只產 markdown",
    )
    render.add_argument("--root", default="/work", help="骨架的根目錄（預設 /work）")
    render.set_defaults(func=cmd_render)

    check = subparsers.add_parser(
        "check", help="列出空欄位、缺 source、模板變數對不到 Minutes Schema"
    )
    check.add_argument("meeting", help="Meeting slug，即 records/ 底下的檔名（不含 .yaml）")
    check.add_argument(
        "--schema",
        required=True,
        help="templates/schema/ 底下的檔名，例如 default.yaml",
    )
    check.add_argument(
        "--markdown-template",
        required=True,
        help="templates/markdown/ 底下的檔名，例如 default.md.j2",
    )
    check.add_argument(
        "--docx-template",
        help="templates/docx/ 底下的檔名，例如 default.docx。不給就只查 markdown",
    )
    check.add_argument("--root", default="/work", help="骨架的根目錄（預設 /work）")
    check.set_defaults(func=cmd_check)

    scan_docx = subparsers.add_parser(
        "scan-docx", help="列出 Docx Source 的段落與表格儲存格供打洞"
    )
    scan_docx.add_argument("source", help="templates/docx-source/ 底下的檔名")
    scan_docx.add_argument("--root", default="/work", help="骨架的根目錄（預設 /work）")
    scan_docx.set_defaults(func=cmd_scan_docx)

    apply_docx = subparsers.add_parser(
        "apply-docx",
        help="依對照表把 Docx Source 打洞成 Docx Template（對照表走 stdin）",
    )
    apply_docx.add_argument("source", help="templates/docx-source/ 底下的檔名")
    apply_docx.add_argument(
        "--output", help="輸出到 templates/docx/ 的檔名，預設與 Docx Source 同名"
    )
    apply_docx.add_argument(
        "--force", action="store_true", help="覆蓋已存在的 Docx Template"
    )
    apply_docx.add_argument("--root", default="/work", help="骨架的根目錄（預設 /work）")
    apply_docx.set_defaults(func=cmd_apply_docx)

    listing = subparsers.add_parser("list", help="回報每個 Meeting 進行到哪一步")
    listing.add_argument("--root", default="/work", help="骨架的根目錄（預設 /work）")
    listing.set_defaults(func=cmd_list)

    return parser


def main(argv=None):
    args = build_parser().parse_args(argv)
    try:
        payload = args.func(args)
    except CommandError as error:
        sys.stderr.write(f"{error}\n")
        return CommandError.EXIT_CODE
    json.dump(payload, sys.stdout, ensure_ascii=False, indent=2)
    sys.stdout.write("\n")
    return 0


if __name__ == "__main__":
    sys.exit(main())
